"""Document text extraction, chunking, embedding and storage.

Supported: PDF, DOCX, TXT, Markdown, CSV, XLSX.
Page numbers are preserved for PDFs; sheet names for XLSX.
"""
from __future__ import annotations

from typing import Iterable

from sqlalchemy.orm import Session

from app.config import settings
from app.models.document import Document, DocumentChunk
from app.services.embedding_service import embed_texts
from app.services.vector_store import get_vector_store
from app.utils.text_utils import clean_text, split_into_chunks


class DocumentProcessingError(Exception):
    pass


def extract_text_pages(file_path: str, file_type: str) -> list[tuple[int | None, str]]:
    """Return a list of (page_number, text). page_number may be None."""
    if file_type == "pdf":
        return _extract_pdf(file_path)
    if file_type == "docx":
        return [(None, _extract_docx(file_path))]
    if file_type in ("txt", "markdown"):
        return [(None, _read_text_file(file_path))]
    if file_type == "csv":
        return [(None, _extract_csv(file_path))]
    if file_type == "xlsx":
        return _extract_xlsx(file_path)
    raise DocumentProcessingError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_path: str) -> list[tuple[int | None, str]]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages: list[tuple[int | None, str]] = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append((idx, text))
    return pages


def _extract_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    # include table text
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_csv(file_path: str) -> str:
    import pandas as pd

    df = pd.read_csv(file_path, dtype=str, keep_default_na=False)
    return _df_to_text(df)


def _extract_xlsx(file_path: str) -> list[tuple[int | None, str]]:
    import pandas as pd

    xls = pd.read_excel(file_path, sheet_name=None, dtype=str)
    out: list[tuple[int | None, str]] = []
    for sheet_name, df in xls.items():
        df = df.fillna("")
        out.append((None, f"# Sheet: {sheet_name}\n{_df_to_text(df)}"))
    return out


def _df_to_text(df) -> str:
    header = " | ".join(str(c) for c in df.columns)
    lines = [header]
    for _, row in df.iterrows():
        lines.append(" | ".join(str(v) for v in row.tolist()))
    return "\n".join(lines)


def process_document(db: Session, document: Document) -> Document:
    """Extract -> chunk -> embed -> store. Updates the document row in place."""
    document.status = "processing"
    db.commit()

    try:
        pages = extract_text_pages(document.file_path, document.file_type)
        total_text = "".join(t for _, t in pages).strip()
        if not total_text:
            raise DocumentProcessingError("The document appears to be empty (no extractable text).")

        chunk_rows: list[DocumentChunk] = []
        vector_ids: list[str] = []
        vector_texts: list[str] = []
        vector_metas: list[dict] = []

        chunk_index = 0
        for page_number, page_text in pages:
            page_text = clean_text(page_text)
            if not page_text:
                continue
            for chunk in split_into_chunks(
                page_text, settings.chunk_size, settings.chunk_overlap
            ):
                vector_id = f"{document.id}:{chunk_index}"
                meta = {
                    "document_id": document.id,
                    "document_name": document.filename,
                    "chunk_index": chunk_index,
                }
                if page_number is not None:
                    meta["page"] = page_number
                row = DocumentChunk(
                    document_id=document.id,
                    chunk_index=chunk_index,
                    content=chunk,
                    meta=meta,
                    vector_id=vector_id,
                )
                chunk_rows.append(row)
                vector_ids.append(vector_id)
                vector_texts.append(chunk)
                vector_metas.append(meta)
                chunk_index += 1

        if not chunk_rows:
            raise DocumentProcessingError("No usable text chunks were produced.")

        embeddings = embed_texts(vector_texts)

        store = get_vector_store()
        store.add(
            ids=vector_ids,
            embeddings=embeddings,
            documents=vector_texts,
            metadatas=vector_metas,
        )

        db.add_all(chunk_rows)
        document.chunk_count = len(chunk_rows)
        document.status = "ready"
        document.error = None
        db.commit()
        db.refresh(document)
        return document

    except DocumentProcessingError as exc:
        document.status = "failed"
        document.error = str(exc)
        db.commit()
        raise
    except Exception as exc:  # pragma: no cover
        document.status = "failed"
        document.error = f"Processing failed: {exc}"
        db.commit()
        raise DocumentProcessingError(str(exc)) from exc


def delete_document_vectors(document_id: str) -> None:
    store = get_vector_store()
    try:
        store.delete(where={"document_id": document_id})
    except Exception as exc:  # pragma: no cover
        print(f"[document_processor] failed to delete vectors for {document_id}: {exc}")

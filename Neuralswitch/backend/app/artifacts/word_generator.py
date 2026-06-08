from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from app.schemas.chat import ChartData, TableData


def generate_word(path: Path, *, title: str, table: TableData, chart: ChartData | None, answer: str | None = None) -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
    if answer:
        doc.add_paragraph(answer)
    doc.add_paragraph(f"Rows: {len(table.rows):,} | Columns: {len(table.columns):,}")
    if chart:
        doc.add_paragraph(f"Suggested chart: {chart.type} chart, {chart.y} by {chart.x}.")

    doc.add_heading("Data", level=2)
    rows = table.rows[:35]
    word_table = doc.add_table(rows=1, cols=max(1, len(table.columns)))
    word_table.style = "Table Grid"
    for index, column in enumerate(table.columns):
        cell = word_table.rows[0].cells[index]
        cell.text = str(column)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = word_table.add_row().cells
        for index, value in enumerate(row[: len(table.columns)]):
            cells[index].text = str(value if value is not None else "")
    if len(table.rows) > len(rows):
        doc.add_paragraph(f"Showing first {len(rows):,} rows. Export Excel/CSV for complete data.")
    doc.add_heading("Recommendations", level=2)
    doc.add_paragraph("Use NeuralSwitch follow-up prompts to drill down by customer, BDM, Geo Head, Practice Head, MS/PS, company, or month.")
    doc.save(path)
